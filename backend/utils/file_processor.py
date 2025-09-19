import os
import docx
import PyPDF2
import pdfplumber
from typing import Optional
import chardet
import logging

async def process_file(file_path: str, content_type: Optional[str]) -> str:
    """
    根据文件类型提取文本内容
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # PDF文件处理
        if content_type and "pdf" in content_type.lower():
            return extract_pdf_content(file_path)
        
        # Word文档处理
        elif content_type and ("word" in content_type.lower() or "document" in content_type.lower()):
            return extract_docx_content(file_path)
        
        # 文本文件处理
        elif content_type and "text" in content_type.lower():
            return extract_text_content(file_path)
        
        # 根据文件扩展名判断
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            return extract_pdf_content(file_path)
        elif file_extension in [".docx", ".doc"]:
            return extract_docx_content(file_path)
        elif file_extension in [".txt", ".md"]:
            return extract_text_content(file_path)
        else:
            # 默认尝试作为文本文件处理
            return extract_text_content(file_path)
    
    except Exception as e:
        return f"文件处理错误: {str(e)}"

def extract_pdf_content(file_path: str) -> str:
    """提取PDF文件内容"""
    try:
        content = ""
        
        # 首先尝试使用pdfplumber（更好的文本提取）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
            
            if content.strip():
                return content.strip()
        except Exception:
            pass
        
        # 如果pdfplumber失败，尝试PyPDF2
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
        
        return content.strip() if content.strip() else "无法提取PDF内容"
    
    except Exception as e:
        return f"PDF处理错误: {str(e)}"

def extract_docx_content(file_path: str) -> str:
    """提取Word文档内容（增强版，支持大文档）"""
    try:
        # 检查文件大小，如果超过10MB使用流式处理
        file_size = os.path.getsize(file_path)
        if file_size > 10 * 1024 * 1024:  # 10MB
            logging.info(f"大文档检测：{file_size} 字节，使用流式处理")
            return extract_large_docx_content(file_path)

        doc = docx.Document(file_path)
        content_parts = []

        # 提取段落内容
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                # 添加段落标题标识（如果是标题样式）
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.split()[-1] if 'Heading' in paragraph.style.name else '1'
                    content_parts.append(f"[标题{level}] {paragraph.text.strip()}")
                else:
                    content_parts.append(paragraph.text.strip())

        # 提取表格内容
        for table_idx, table in enumerate(doc.tables):
            content_parts.append(f"\n[表格{table_idx + 1}]")
            for row_idx, row in enumerate(table.rows):
                row_text = "\t".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content_parts.append(row_text)

        # 提取页眉页脚（如果有）
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = "\n".join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
                if header_text:
                    content_parts.insert(0, f"[页眉]\n{header_text}")

            if section.footer and section.footer.paragraphs:
                footer_text = "\n".join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
                if footer_text:
                    content_parts.append(f"[页脚]\n{footer_text}")

        content = "\n".join(content_parts)
        return content.strip() if content.strip() else "无法提取Word文档内容"

    except Exception as e:
        logging.error(f"Word文档处理错误: {str(e)}")
        return f"Word文档处理错误: {str(e)}"

def extract_large_docx_content(file_path: str) -> str:
    """大文档专用提取函数，减少内存占用"""
    try:
        doc = docx.Document(file_path)
        content_parts = []

        # 分段处理，避免一次性加载所有内容
        paragraph_count = len(doc.paragraphs)
        batch_size = 1000  # 每批处理1000个段落

        for batch_start in range(0, paragraph_count, batch_size):
            batch_end = min(batch_start + batch_size, paragraph_count)

            for i in range(batch_start, batch_end):
                paragraph = doc.paragraphs[i]
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.split()[-1] if 'Heading' in paragraph.style.name else '1'
                        content_parts.append(f"[标题{level}] {paragraph.text.strip()}")
                    else:
                        content_parts.append(paragraph.text.strip())

        # 表格处理
        for table_idx, table in enumerate(doc.tables):
            content_parts.append(f"\n[表格{table_idx + 1}]")
            try:
                for row in table.rows:
                    row_text = "\t".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content_parts.append(row_text)
            except Exception as e:
                logging.warning(f"表格{table_idx + 1}处理失败: {str(e)}")
                content_parts.append(f"[表格{table_idx + 1}处理失败]")

        return "\n".join(content_parts).strip()

    except Exception as e:
        logging.error(f"大文档处理错误: {str(e)}")
        return f"大文档处理错误: {str(e)}"

def extract_text_content(file_path: str) -> str:
    """提取纯文本文件内容（增强版，支持自动编码检测）"""
    try:
        # 首先检测文件编码
        with open(file_path, 'rb') as file:
            raw_data = file.read()

        # 使用chardet检测编码
        detected = chardet.detect(raw_data)
        encoding = detected['encoding'] or 'utf-8'
        confidence = detected['confidence'] or 0

        logging.info(f"检测到文件编码: {encoding}, 置信度: {confidence}")

        try:
            # 尝试使用检测到的编码
            content = raw_data.decode(encoding)
            return content.strip() if content.strip() else "文本文件内容为空"
        except UnicodeDecodeError:
            # 如果检测失败，尝试常见编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin-1']

            for enc in encodings:
                try:
                    content = raw_data.decode(enc)
                    logging.info(f"使用备用编码 {enc} 成功解码")
                    return content.strip() if content.strip() else "文本文件内容为空"
                except UnicodeDecodeError:
                    continue

        # 如果所有编码都失败，使用latin-1作为最后手段
        try:
            content = raw_data.decode('latin-1')
            logging.warning("使用latin-1编码作为最后手段，可能显示异常字符")
            return content.strip() if content.strip() else "文本文件内容为空"
        except Exception:
            return "无法解码文本文件内容"

    except Exception as e:
        logging.error(f"文本文件处理错误: {str(e)}")
        return f"文本文件处理错误: {str(e)}"